from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
REPO_URL = "https://github.com/Dew3120/huddle.git"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF6F1"
PALE_WARNING = "FFF4E5"
BORDER = "CDD7E1"
CODE_FILL = "F5F7FA"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_keep_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_doc(title, member):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        footer.add_run(f"Huddle Assignment 01 | {member} contribution guide"),
        size=8,
        color=MUTED,
    )
    return doc


def add_title_block(doc, member, feature, branch, order_note):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("HUDDLE ASSIGNMENT 01"), size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(f"{member}'s Contribution Guide"), size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(feature), size=13, color=DARK_BLUE, bold=True)

    table = doc.add_table(rows=5, cols=2)
    set_table_width(table, [2000, 7360])
    rows = [
        ("Repository", REPO_URL),
        ("Branch", branch),
        ("Base branch", "main"),
        ("Important order", order_note),
        ("Rule", "Do not commit to main, do not force-push, and do not merge your own pull request."),
    ]
    for row_idx, (label, value) in enumerate(rows):
        for col_idx, text in enumerate((label, value)):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col_idx == 0:
                shade_cell(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5, color=INK, bold=(col_idx == 0))

    add_callout(
        doc,
        "How to use this file",
        "Follow the commands and code in order. If an error appears, stop and send the full screenshot or terminal text to the team leader. Do not delete files or guess a fix.",
        PALE_WARNING,
    )


def h1(doc, text):
    paragraph = doc.add_heading(text, level=1)
    add_keep_next(paragraph)


def h2(doc, text):
    paragraph = doc.add_heading(text, level=2)
    add_keep_next(paragraph)


def h3(doc, text):
    paragraph = doc.add_heading(text, level=3)
    add_keep_next(paragraph)


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[ ] {item}")


def code_block(doc, text, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(caption), size=10, color=BLUE, bold=True)

    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    set_cell_border(cell, color="D7DEE8", size="4")
    shade_cell(cell, CODE_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(text.strip("\n").splitlines()):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=7.4, color="1F2937")


def add_callout(doc, label, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    set_cell_border(cell, color=BORDER, size="4")
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10, color=DARK_BLUE, bold=True)
    set_run_font(paragraph.add_run(text), size=10, color=INK)


def add_setup_steps(doc, member, branch, dependency):
    h1(doc, "1. Set up Git safely")
    add_callout(doc, "Required order", dependency, PALE_GREEN)

    h2(doc, "1.1 Open or clone the repository")
    para(doc, "Use the first block if the repository already exists on your computer. Use the second block only if it does not exist yet.")
    code_block(
        doc,
        r"""cd C:\Users\$env:USERNAME\huddle
git status --short""",
        "Existing local repository",
    )
    code_block(
        doc,
        rf"""cd C:\Users\$env:USERNAME
git clone {REPO_URL}
cd huddle""",
        "Clone if needed",
    )

    h2(doc, "1.2 Set your own Git identity")
    para(doc, "Replace the email with the email connected to your GitHub account. This makes your commit appear as your own contribution.")
    code_block(
        doc,
        rf"""git config user.name "{member}"
git config user.email "your-github-email@example.com"
git config user.name
git config user.email""",
    )
    add_callout(doc, "Expected", f"The last two lines must show {member}'s name and email, not the leader's details.", PALE_GREEN)

    h2(doc, "1.3 Update main and create your feature branch")
    code_block(
        doc,
        f"""git checkout main
git pull origin main
git checkout -b {branch}
git branch --show-current""",
    )
    add_callout(doc, "Expected", f"The final command prints {branch}. If the branch already exists, run git checkout {branch}.", PALE_GREEN)

    h2(doc, "1.4 Install dependencies and open the app")
    code_block(
        doc,
        """npm install
npm run dev""",
    )
    para(doc, "Open the Local URL printed by Vite, usually http://localhost:5173/. Keep this terminal running while testing. Stop it with Ctrl+C before committing.")


def add_finish_steps(doc, branch, add_paths, commit_title, commit_body, pr_title, pr_body, tests, explanation_points):
    h1(doc, "5. Test before committing")
    checklist(doc, tests)
    para(doc, "After browser testing passes, stop the dev server with Ctrl+C and run these checks.")
    code_block(
        doc,
        """npm run build
git diff --check
git status --short""",
    )
    add_callout(doc, "Expected", "The build succeeds, git diff --check prints nothing, and git status shows only your intended files.", PALE_GREEN)

    h1(doc, "6. Commit and push")
    para(doc, "Do not use git add . here. Add only the files for your assigned feature.")
    code_block(
        doc,
        f"""git add {add_paths}
git commit -m "{commit_title}" -m "{commit_body}"
git push -u origin {branch}
git status --short""",
    )
    add_callout(doc, "Expected", "The push succeeds and the final git status is clean.", PALE_GREEN)

    h1(doc, "7. Open the pull request")
    numbers(
        doc,
        [
            "Open the Huddle repository on GitHub.",
            f"Create a pull request from {branch} into main.",
            "Paste the title and description below.",
            "Request the team leader as reviewer.",
            "Do not merge your own pull request.",
        ],
    )
    code_block(doc, pr_title, "Pull request title")
    code_block(doc, pr_body, "Pull request description")

    h1(doc, "8. Be ready to explain your work")
    para(doc, "Before the pull request is accepted, explain these points in your own words.")
    bullets(doc, explanation_points)

    h1(doc, "9. Final checklist")
    checklist(
        doc,
        [
            "I used my own Git name and verified GitHub email.",
            "I worked on my assigned feature branch, not main.",
            "I typed/read the code and can explain it.",
            "The feature works in the browser.",
            "npm run build passed.",
            "git diff --check printed no errors.",
            "My commit was pushed to GitHub from my account.",
            "I opened a pull request into main.",
            "I did not merge my own pull request.",
        ],
    )


BUTTON_JSX = """import styles from './Button.module.css';

export default function Button({
  variant = 'primary',
  size = 'medium',
  type = 'button',
  className = '',
  children,
  ...props
}) {
  const classNames = [
    styles.button,
    styles[variant],
    styles[size],
    className,
  ]
    .filter(Boolean)
    .join(' ');

  return (
    <button type={type} className={classNames} {...props}>
      {children}
    </button>
  );
}
"""


BUTTON_CSS = """.button {
  display: inline-flex;
  min-height: 42px;
  align-items: center;
  justify-content: center;
  gap: 8px;
  padding: 9px 18px;
  border: 1px solid transparent;
  border-radius: 6px;
  font: inherit;
  font-weight: 700;
  line-height: 1;
  cursor: pointer;
  transition:
    background-color 150ms ease,
    border-color 150ms ease,
    color 150ms ease;
}

.button:focus-visible {
  outline: 3px solid rgb(53 120 168 / 25%);
  outline-offset: 2px;
}

.button:disabled {
  cursor: not-allowed;
  opacity: 0.45;
}

.primary {
  border-color: #176b55;
  background: #176b55;
  color: #ffffff;
}

.primary:hover:not(:disabled) {
  border-color: #125744;
  background: #125744;
}

.secondary {
  border-color: #b8c5d4;
  background: #ffffff;
  color: #1e3a5f;
}

.secondary:hover:not(:disabled) {
  border-color: #3578a8;
  background: #edf5fb;
}

.danger {
  border-color: #e7b6b6;
  background: #ffffff;
  color: #a51d2d;
}

.danger:hover:not(:disabled) {
  border-color: #d99090;
  background: #fff1f1;
}

.small {
  min-height: 36px;
  padding: 7px 10px;
  font-size: 0.8rem;
}

.medium {
  min-height: 42px;
  padding: 9px 18px;
}
"""


ADD_TASK_FORM_FINAL = """import { useState } from 'react';
import Button from './Button/Button.jsx';

const initialForm = {
  title: '',
  assignee: '',
  dueDate: '',
};

function getToday() {
  const date = new Date();
  const year = date.getFullYear();
  const month = String(date.getMonth() + 1).padStart(2, '0');
  const day = String(date.getDate()).padStart(2, '0');

  return `${year}-${month}-${day}`;
}

export default function AddTaskForm({ onAdd }) {
  const [form, setForm] = useState(initialForm);
  const [errors, setErrors] = useState({});

  const today = getToday();

  function handleChange(event) {
    const { name, value } = event.target;

    setForm((current) => ({
      ...current,
      [name]: value,
    }));

    setErrors((current) => ({
      ...current,
      [name]: '',
    }));
  }

  function validateForm() {
    const nextErrors = {};
    const title = form.title.trim();
    const assignee = form.assignee.trim();

    if (!title) {
      nextErrors.title = 'Enter a task title.';
    } else if (title.length < 3) {
      nextErrors.title = 'The title must contain at least 3 characters.';
    }

    if (!assignee) {
      nextErrors.assignee = 'Enter the name of an assignee.';
    }

    if (!form.dueDate) {
      nextErrors.dueDate = 'Select a due date.';
    } else if (form.dueDate < today) {
      nextErrors.dueDate = 'The due date cannot be in the past.';
    }

    return nextErrors;
  }

  function handleSubmit(event) {
    event.preventDefault();

    const nextErrors = validateForm();

    if (Object.keys(nextErrors).length > 0) {
      setErrors(nextErrors);
      return;
    }

    onAdd({
      title: form.title.trim(),
      assignee: form.assignee.trim(),
      dueDate: form.dueDate,
    });

    setForm(initialForm);
    setErrors({});
  }

  return (
    <form className="task-form" onSubmit={handleSubmit} noValidate>
      <label>
        Task title
        <input
          name="title"
          value={form.title}
          onChange={handleChange}
          aria-invalid={Boolean(errors.title)}
          aria-describedby={errors.title ? 'title-error' : undefined}
        />
        {errors.title && (
          <span className="form-error" id="title-error">
            {errors.title}
          </span>
        )}
      </label>

      <label>
        Assignee
        <input
          name="assignee"
          value={form.assignee}
          onChange={handleChange}
          aria-invalid={Boolean(errors.assignee)}
          aria-describedby={
            errors.assignee ? 'assignee-error' : undefined
          }
        />
        {errors.assignee && (
          <span className="form-error" id="assignee-error">
            {errors.assignee}
          </span>
        )}
      </label>

      <label>
        Due date
        <input
          type="date"
          name="dueDate"
          value={form.dueDate}
          min={today}
          onChange={handleChange}
          aria-invalid={Boolean(errors.dueDate)}
          aria-describedby={
            errors.dueDate ? 'due-date-error' : undefined
          }
        />
        {errors.dueDate && (
          <span className="form-error" id="due-date-error">
            {errors.dueDate}
          </span>
        )}
      </label>

      <Button type="submit">Add task</Button>
    </form>
  );
}
"""


EDIT_TASK_REPLACEMENT = """import Button from './Button/Button.jsx';

...

<div className="task-edit-form__actions">
  <Button type="button" variant="secondary" onClick={onCancel}>
    Cancel
  </Button>
  <Button type="submit">Save changes</Button>
</div>
"""


TASK_CARD_FINAL = """import { Link } from 'react-router-dom';
import Button from './Button/Button.jsx';

const statusLabels = {
  todo: 'To Do',
  'in-progress': 'In Progress',
  done: 'Done',
};

export default function TaskCard({
  id,
  title,
  assignee = 'Unassigned',
  status = 'todo',
  dueDate,
  onDelete,
  onMove,
}) {
  const cannotMoveLeft = status === 'todo';
  const cannotMoveRight = status === 'done';

  return (
    <article className={`task-card task-card--${status}`}>
      <span className="task-card__status">
        {statusLabels[status] ?? 'Unknown'}
      </span>

      <h3 className="task-card__title">
        <Link className="task-card__link" to={`/tasks/${id}`}>
          {title}
        </Link>
      </h3>

      <dl className="task-card__details">
        <div>
          <dt>Assignee</dt>
          <dd>{assignee}</dd>
        </div>

        <div>
          <dt>Due date</dt>
          <dd>
            <time dateTime={dueDate}>{dueDate}</time>
          </dd>
        </div>
      </dl>

      <div
        className="task-card__actions"
        aria-label={`Actions for ${title}`}
      >
        <Button
          type="button"
          variant="secondary"
          size="small"
          onClick={() => onMove(id, 'left')}
          disabled={cannotMoveLeft}
          aria-label={`Move ${title} left`}
        >
          Move left
        </Button>

        <Button
          type="button"
          variant="secondary"
          size="small"
          onClick={() => onMove(id, 'right')}
          disabled={cannotMoveRight}
          aria-label={`Move ${title} right`}
        >
          Move right
        </Button>

        <Button
          type="button"
          variant="danger"
          size="small"
          className="task-card__delete"
          onClick={() => onDelete(id)}
          aria-label={`Delete ${title}`}
        >
          Delete
        </Button>
      </div>
    </article>
  );
}
"""


ERROR_STATE_FINAL = """import Button from './Button/Button.jsx';

export default function ErrorState({ message, onRetry }) {
  return (
    <section className="screen-state screen-state--error" role="alert">
      <h2>Tasks could not be loaded</h2>
      <p>{message}</p>
      <Button type="button" variant="secondary" onClick={onRetry}>
        Try again
      </Button>
    </section>
  );
}
"""


TASK_DETAIL_REPLACEMENT = """import Button from '../components/Button/Button.jsx';

...

{!isEditing && (
  <Button type="button" onClick={() => setIsEditing(true)}>
    Edit task
  </Button>
)}
"""


FILTER_TASKS_JS = """export function getAssignees(tasks) {
  return [...new Set(tasks.map((task) => task.assignee).filter(Boolean))].sort(
    (first, second) => first.localeCompare(second),
  );
}

export function isOverdue(task, today = new Date()) {
  if (!task.dueDate || task.status === 'done') {
    return false;
  }

  const dueDate = new Date(`${task.dueDate}T23:59:59`);

  return !Number.isNaN(dueDate.getTime()) && dueDate < today;
}

export function filterTasks(tasks, filters) {
  const query = filters.query.trim().toLowerCase();

  return tasks.filter((task) => {
    const matchesQuery =
      query.length === 0 || task.title.toLowerCase().includes(query);
    const matchesAssignee =
      filters.assignee.length === 0 || task.assignee === filters.assignee;
    const matchesStatus =
      filters.status.length === 0 || task.status === filters.status;
    const matchesOverdue = !filters.overdue || isOverdue(task);

    return (
      matchesQuery &&
      matchesAssignee &&
      matchesStatus &&
      matchesOverdue
    );
  });
}
"""


TASK_FILTERS_JSX = """import Button from './Button/Button.jsx';

export default function TaskFilters({
  filters,
  assignees,
  resultCount,
  totalCount,
  onChange,
  onClear,
}) {
  return (
    <section className="task-filters" aria-label="Task filters">
      <label>
        Search title
        <input
          type="search"
          value={filters.query}
          onChange={(event) => onChange('query', event.target.value)}
          placeholder="Search task titles"
        />
      </label>

      <label>
        Assignee
        <select
          value={filters.assignee}
          onChange={(event) => onChange('assignee', event.target.value)}
        >
          <option value="">All assignees</option>
          {assignees.map((assignee) => (
            <option key={assignee} value={assignee}>
              {assignee}
            </option>
          ))}
        </select>
      </label>

      <label>
        Status
        <select
          value={filters.status}
          onChange={(event) => onChange('status', event.target.value)}
        >
          <option value="">All statuses</option>
          <option value="todo">To Do</option>
          <option value="in-progress">In Progress</option>
          <option value="done">Done</option>
        </select>
      </label>

      <label className="task-filters__checkbox">
        <input
          type="checkbox"
          checked={filters.overdue}
          onChange={(event) => onChange('overdue', event.target.checked)}
        />
        Overdue only
      </label>

      <Button type="button" variant="secondary" onClick={onClear}>
        Clear filters
      </Button>

      <p className="task-filters__summary" aria-live="polite">
        Showing {resultCount} of {totalCount} tasks
      </p>
    </section>
  );
}
"""


BOARD_PAGE_FINAL = """import { useMemo } from 'react';
import { useSearchParams } from 'react-router-dom';
import AddTaskForm from '../components/AddTaskForm.jsx';
import Board from '../components/Board.jsx';
import Button from '../components/Button/Button.jsx';
import ErrorState from '../components/ErrorState.jsx';
import LoadingState from '../components/LoadingState.jsx';
import TaskFilters from '../components/TaskFilters.jsx';
import { useTasks } from '../hooks/useTasks.js';
import { filterTasks, getAssignees } from '../utils/filterTasks.js';

export default function BoardPage() {
  const [searchParams, setSearchParams] = useSearchParams();
  const {
    tasks,
    loading,
    error,
    addTask,
    deleteTask,
    moveTask,
    retryLoading,
  } = useTasks();

  const filters = {
    query: searchParams.get('q') ?? '',
    assignee: searchParams.get('assignee') ?? '',
    status: searchParams.get('status') ?? '',
    overdue: searchParams.get('overdue') === 'true',
  };

  const assignees = useMemo(() => getAssignees(tasks), [tasks]);
  const visibleTasks = useMemo(
    () => filterTasks(tasks, filters),
    [tasks, filters.query, filters.assignee, filters.status, filters.overdue],
  );
  const completedCount = tasks.filter(
    (task) => task.status === 'done',
  ).length;
  const hasActiveFilters =
    filters.query ||
    filters.assignee ||
    filters.status ||
    filters.overdue;

  function handleDelete(taskId) {
    if (window.confirm('Delete this task permanently?')) {
      deleteTask(taskId);
    }
  }

  function updateFilter(name, value) {
    const nextParams = new URLSearchParams(searchParams);
    const key = name === 'query' ? 'q' : name;

    if (value === '' || value === false) {
      nextParams.delete(key);
    } else {
      nextParams.set(key, String(value));
    }

    setSearchParams(nextParams, { replace: true });
  }

  function clearFilters() {
    setSearchParams({}, { replace: true });
  }

  return (
    <main className="app-shell">
      <header className="page-header">
        <div>
          <p>Huddle workspace</p>
          <h1>Team Task Board</h1>
          <span>Plan, assign, and track your team&apos;s work.</span>
        </div>

        {!loading && !error && (
          <strong className="board-progress" aria-live="polite">
            {completedCount} of {tasks.length} done
          </strong>
        )}
      </header>

      {loading ? (
        <LoadingState />
      ) : error ? (
        <ErrorState message={error} onRetry={retryLoading} />
      ) : (
        <>
          <AddTaskForm onAdd={addTask} />
          <TaskFilters
            filters={filters}
            assignees={assignees}
            resultCount={visibleTasks.length}
            totalCount={tasks.length}
            onChange={updateFilter}
            onClear={clearFilters}
          />

          {visibleTasks.length > 0 ? (
            <Board
              tasks={visibleTasks}
              onDelete={handleDelete}
              onMove={moveTask}
            />
          ) : (
            <section className="empty-state" role="status">
              <h2>No matching tasks</h2>
              <p>
                {hasActiveFilters
                  ? 'Change or clear the filters to see tasks again.'
                  : 'Create a task to start the board.'}
              </p>
              {hasActiveFilters && (
                <Button
                  type="button"
                  variant="secondary"
                  onClick={clearFilters}
                >
                  Clear filters
                </Button>
              )}
            </section>
          )}
        </>
      )}
    </main>
  );
}
"""


FILTER_CSS = """.task-filters {
  display: grid;
  grid-template-columns: minmax(220px, 2fr) minmax(150px, 1fr) minmax(150px, 1fr) auto auto;
  align-items: end;
  gap: 14px;
  max-width: 1400px;
  margin: 0 auto 24px;
  padding: 18px;
  border: 1px solid #d7e0ea;
  border-radius: 8px;
  background: #ffffff;
}

.task-filters label {
  display: grid;
  gap: 7px;
  color: #334155;
  font-size: 0.875rem;
  font-weight: 600;
}

.task-filters input[type='search'],
.task-filters select {
  width: 100%;
  min-height: 42px;
  padding: 9px 11px;
  border: 1px solid #cbd5e1;
  border-radius: 6px;
  background: #ffffff;
  color: #0f172a;
  font: inherit;
}

.task-filters__checkbox {
  display: flex !important;
  min-height: 42px;
  align-items: center;
  gap: 8px;
  white-space: nowrap;
}

.task-filters__checkbox input {
  width: 18px;
  min-height: 18px;
  margin: 0;
}

.task-filters__summary {
  grid-column: 1 / -1;
  margin: 0;
  color: #667085;
  font-size: 0.875rem;
  font-weight: 600;
}

.empty-state {
  display: grid;
  justify-items: center;
  gap: 8px;
  max-width: 1400px;
  margin: 0 auto;
  padding: 40px 24px;
  border: 1px dashed #b8c5d4;
  border-radius: 8px;
  background: #ffffff;
  text-align: center;
}

.empty-state h2,
.empty-state p {
  margin: 0;
}

.empty-state p {
  color: #667085;
}

@media (max-width: 1000px) {
  .task-filters {
    grid-template-columns: 1fr 1fr;
  }
}

@media (max-width: 600px) {
  .task-filters {
    grid-template-columns: 1fr;
  }
}
"""


def create_charles_guide():
    member = "Charles"
    branch = "feature/shared-button"
    doc = configure_doc("Charles shared button guide", member)
    add_title_block(
        doc,
        member,
        "Contribution: shared reusable Button component and consistent button usage",
        branch,
        "Start only after the leader has merged feature/frontend-assignment into main.",
    )

    h1(doc, "Your exact outcome")
    bullets(
        doc,
        [
            "Create src/components/Button/Button.jsx.",
            "Create src/components/Button/Button.module.css.",
            "Replace raw button elements in the current React files with the shared Button component.",
            "Keep form submit behavior, click handlers, disabled states, aria-labels, and keyboard focus working.",
            "Commit and push from Charles's own Git account.",
        ],
    )

    add_setup_steps(
        doc,
        member,
        branch,
        "The leader's frontend foundation branch must already be merged into main before Charles starts.",
    )

    h1(doc, "2. Create the Button component")
    h2(doc, "2.1 Make the folder")
    code_block(doc, r"mkdir src\components\Button")
    h2(doc, "2.2 Create src/components/Button/Button.jsx")
    code_block(doc, BUTTON_JSX, "Paste the complete file")
    h2(doc, "2.3 Create src/components/Button/Button.module.css")
    code_block(doc, BUTTON_CSS, "Paste the complete file")

    h1(doc, "3. Replace raw buttons")
    para(doc, "Run this command first so you can see every raw button that still needs attention.")
    code_block(doc, r'rg -n "<button|</button>" .\src')

    h2(doc, "3.1 Replace src/components/AddTaskForm.jsx")
    para(doc, "This is the complete final file for AddTaskForm.jsx after adding Button.")
    code_block(doc, ADD_TASK_FORM_FINAL)

    h2(doc, "3.2 Edit src/components/EditTaskForm.jsx")
    para(doc, "Add the import at the top, then replace only the action buttons at the bottom. Leave validation and state code unchanged.")
    code_block(doc, EDIT_TASK_REPLACEMENT)

    h2(doc, "3.3 Replace src/components/TaskCard.jsx")
    para(doc, "This complete version preserves the current prop names: id, title, status, onMove, and onDelete.")
    code_block(doc, TASK_CARD_FINAL)

    h2(doc, "3.4 Replace src/components/ErrorState.jsx")
    code_block(doc, ERROR_STATE_FINAL)

    h2(doc, "3.5 Edit src/pages/TaskDetailPage.jsx")
    para(doc, "Add the import and replace the Edit task raw button. Leave the rest of the file unchanged.")
    code_block(doc, TASK_DETAIL_REPLACEMENT)

    h2(doc, "3.6 Confirm only the shared component has a raw button")
    code_block(doc, r'rg -n "<button|</button>" .\src')
    add_callout(
        doc,
        "Expected",
        "The only raw button result should be src/components/Button/Button.jsx. If another file appears, convert that button while preserving its existing props and handlers.",
        PALE_GREEN,
    )

    h1(doc, "4. Manual browser test")
    para(doc, "Run npm run dev, open the local URL, and test every interaction below.")

    add_finish_steps(
        doc,
        branch,
        "src/components/Button src/components/AddTaskForm.jsx src/components/EditTaskForm.jsx src/components/TaskCard.jsx src/components/ErrorState.jsx src/pages/TaskDetailPage.jsx",
        "Add reusable button component",
        "Create a shared Button with primary, secondary, and danger variants and replace repeated raw buttons without changing behavior.",
        "Add shared reusable Button component",
        """## Summary
- Added a reusable Button component.
- Added primary, secondary, danger, medium, and small button styles.
- Replaced repeated raw buttons while preserving click, submit, disabled, and aria behavior.

## Verification
- npm run build
- git diff --check
- Manual browser test: add, edit, cancel, move, delete, retry, edit-task, disabled buttons, and keyboard focus.

## Assignment evidence
This is Charles's individual Assignment 01 contribution.""",
        [
            "Create a valid task with Add task.",
            "Open a task detail page and click Edit task.",
            "Use Save changes and confirm edits save.",
            "Use Cancel and confirm it exits edit mode without saving.",
            "Move a task left and right.",
            "Confirm Move left is disabled in To Do and Move right is disabled in Done.",
            "Delete a task and confirm the existing browser confirmation still appears.",
            "Trigger or inspect the error retry button if possible.",
            "Use Tab to focus buttons and confirm a visible focus outline appears.",
            "Open DevTools Console and confirm there are no red errors.",
        ],
        [
            "The Button component renders one real HTML button and forwards native props through ...props.",
            "variant controls the visual intent: primary, secondary, or danger.",
            "size controls the button dimensions without changing behavior.",
            "children is the visible text inside the button.",
            "CSS Modules keep Button styles scoped so they do not accidentally style unrelated elements.",
            "Replacing a raw button is safe only when the original type, onClick, disabled, and aria props are preserved.",
        ],
    )

    path = OUTPUT_DIR / "Charles-Assignment-01-Guide.docx"
    doc.save(path)
    return path


def create_vinuka_guide():
    member = "Vinuka"
    branch = "feature/task-filters"
    doc = configure_doc("Vinuka task filters guide", member)
    add_title_block(
        doc,
        member,
        "Contribution: task search, filters, URL state, and no-results handling",
        branch,
        "Start only after Charles's shared Button pull request has been merged into main.",
    )

    h1(doc, "Your exact outcome")
    bullets(
        doc,
        [
            "Create src/utils/filterTasks.js.",
            "Create src/components/TaskFilters.jsx.",
            "Update src/pages/BoardPage.jsx so filters read from and write to the URL.",
            "Append filter and no-results styles to src/index.css.",
            "Keep add, move, delete, loading, error, and progress behavior working.",
            "Commit and push from Vinuka's own Git account.",
        ],
    )

    add_setup_steps(
        doc,
        member,
        branch,
        "Vinuka starts after Charles's Button component is merged because TaskFilters imports that shared Button.",
    )

    h1(doc, "2. Create the filter utility")
    h2(doc, "2.1 Create src/utils/filterTasks.js")
    code_block(doc, FILTER_TASKS_JS, "Paste the complete file")
    add_callout(
        doc,
        "Why this file exists",
        "The filtering logic stays separate from the UI, which makes BoardPage easier to read and makes this logic easier to unit test later.",
        LIGHT_BLUE,
    )

    h1(doc, "3. Create the filter component")
    h2(doc, "3.1 Create src/components/TaskFilters.jsx")
    code_block(doc, TASK_FILTERS_JSX, "Paste the complete file")

    h1(doc, "4. Wire filters into the board page")
    h2(doc, "4.1 Replace src/pages/BoardPage.jsx")
    para(doc, "This is the complete final BoardPage.jsx. It keeps the original loading, error, add, move, delete, and progress behavior.")
    code_block(doc, BOARD_PAGE_FINAL)

    h2(doc, "4.2 Append styles to src/index.css")
    para(doc, "Scroll to the end of src/index.css and paste this below the existing styles.")
    code_block(doc, FILTER_CSS)

    add_finish_steps(
        doc,
        branch,
        "src/utils/filterTasks.js src/components/TaskFilters.jsx src/pages/BoardPage.jsx src/index.css",
        "Add task search and filters",
        "Add title search, assignee, status, and overdue filters with URL-backed state and a no-results view.",
        "Add task search and URL-backed filters",
        """## Summary
- Added title search, assignee filter, status filter, and overdue-only filter.
- Stored filter state in the URL query string.
- Added result count, clear filters, and no-results behavior.
- Kept existing task add, move, delete, loading, and error behavior working.

## Verification
- npm run build
- git diff --check
- Manual browser test: search, assignee, status, overdue, combined filters, refresh persistence, clear filters, no-results state, add, move, and delete.

## Assignment evidence
This is Vinuka's individual Assignment 01 contribution.""",
        [
            "Search by part of a task title; matching is case-insensitive.",
            "Choose each assignee and confirm only that person's tasks appear.",
            "Choose To Do, In Progress, and Done and confirm each status filter works.",
            "Enable Overdue only and confirm completed tasks are not treated as overdue.",
            "Combine search, assignee, status, and overdue filters and confirm all conditions apply together.",
            "Confirm the browser URL changes, for example ?q=login&status=todo.",
            "Refresh the browser and confirm the same filters remain selected.",
            "Choose filters that match nothing and confirm the no-results state appears.",
            "Click Clear filters and confirm all tasks return and the URL query string is removed.",
            "Create a task and confirm it still appears in To Do.",
            "Move and delete tasks after filtering and confirm the actions still work.",
            "Resize the browser and confirm the filter controls stack without overlap.",
            "Open DevTools Console and confirm there are no red errors.",
        ],
        [
            "filterTasks does not change task data; it returns a visible subset for display.",
            "URLSearchParams stores filter choices in the query string so refresh and shared links preserve them.",
            "The q URL key maps to the query field in React state.",
            "getAssignees uses a Set to remove duplicate assignee names.",
            "isOverdue excludes completed tasks and checks unfinished due dates against today.",
            "The progress counter still uses all tasks, while Board receives only visibleTasks.",
            "The no-results state means data loaded successfully but the current filters match no tasks.",
        ],
    )

    path = OUTPUT_DIR / "Vinuka-Assignment-01-Guide.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for output in (create_charles_guide(), create_vinuka_guide()):
        print(output)
