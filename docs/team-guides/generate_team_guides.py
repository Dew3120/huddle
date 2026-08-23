from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
REPO_URL = "https://github.com/Dew3120/huddle.git"

NAVY = "102A43"
BLUE = "1F6F8B"
PALE_BLUE = "EAF3F7"
PALE_GREEN = "EAF6F1"
PALE_ORANGE = "FFF3E6"
LIGHT_GREY = "F4F6F8"
MID_GREY = "667085"
WHITE = "FFFFFF"
RED = "B42318"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D0D5DD", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document(title):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("263238")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 28, NAVY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11.5, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    styles["Title"].paragraph_format.space_before = Pt(0)

    for footer in (section.footer,):
        table = footer.add_table(rows=1, cols=2, width=Inches(7.0))
        table.columns[0].width = Inches(5.8)
        table.columns[1].width = Inches(1.2)
        left = table.cell(0, 0).paragraphs[0]
        left_run = left.add_run(f"Huddle Assignment 01 | {title}")
        left_run.font.size = Pt(8)
        left_run.font.color.rgb = RGBColor.from_string(MID_GREY)
        add_page_number(table.cell(0, 1).paragraphs[0])

    return doc


def add_cover(doc, member, role, branch, feature):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    run = p.add_run("HUDDLE")
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Assignment 01\nContribution Guide")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(member)
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.6)
    table.columns[1].width = Inches(4.8)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, 120, 150, 120, 150)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [
        ("Role", role),
        ("Feature", feature),
        ("Branch", branch),
        ("Repository", "github.com/Dew3120/huddle"),
    ]
    for index, (label, value) in enumerate(rows):
        shade_cell(table.cell(index, 0), PALE_BLUE)
        table.cell(index, 0).paragraphs[0].add_run(label).bold = True
        table.cell(index, 1).paragraphs[0].add_run(value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Follow every step in order. Type and understand the code yourself, test it, "
        "commit it using your own Git identity, and ask the team leader to review the pull request."
    )
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(MID_GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for the Huddle team | 19 August 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color=BLUE if fill == PALE_BLUE else "D0A245")
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "102A43")
    set_cell_border(cell, color="102A43")
    set_cell_margins(cell, 120, 140, 120, 140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code.rstrip())
    run.font.name = "Cascadia Mono"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cascadia Mono")
    run.font.size = Pt(7.6)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_command(doc, command, caption=None):
    add_code(doc, command, caption)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.add_run("☐ ").font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(item)


def add_contents(doc, sections):
    add_heading(doc, "How to use this guide", 1)
    doc.add_paragraph(
        "This is a literal start-to-finish handoff. Do not skip the Git, test, or pull-request steps. "
        "The feature is only complete after it works, builds successfully, is pushed from your own account, "
        "and is reviewed by the team leader."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Part"
    headers[1].text = "Outcome"
    set_repeat_table_header(table.rows[0])
    for cell in headers:
        shade_cell(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for number, (name, outcome) in enumerate(sections, 1):
        cells = table.add_row().cells
        cells[0].text = f"{number}. {name}"
        cells[1].text = outcome
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_page_break()


def add_common_setup(doc, member, branch, depends_on=None):
    add_heading(doc, "Before touching the code", 1)
    add_callout(
        doc,
        "Important",
        "Work on your own feature branch. Never commit directly to main. Do not use the leader's Git name or email. "
        "Your contribution must appear under your own GitHub identity.",
        PALE_ORANGE,
    )
    if depends_on:
        add_callout(doc, "Required order", depends_on, PALE_GREEN)

    add_heading(doc, "1. Open PowerShell and locate the repository", 2)
    doc.add_paragraph(
        "If the repository is already on your computer, use the first block. If it is not, use the clone block. "
        "Replace <Windows-user> with the name shown after C:\\Users\\ on your computer."
    )
    add_command(
        doc,
        "cd C:\\Users\\<Windows-user>\\huddle\ngit status --short",
        "Repository already exists",
    )
    add_command(
        doc,
        f"cd C:\\Users\\<Windows-user>\ngit clone {REPO_URL}\ncd huddle",
        "Clone only if the folder does not exist",
    )

    add_heading(doc, "2. Set your own Git identity", 2)
    doc.add_paragraph(
        f"Use {member}'s own GitHub display name and a verified email address. These commands configure only this Huddle repository."
    )
    add_command(
        doc,
        f'git config user.name "{member}"\ngit config user.email "<your-verified-GitHub-email>"\n'
        "git config user.name\ngit config user.email",
    )
    add_callout(
        doc,
        "Expected result",
        f"The final two lines print {member}'s name and email. Stop and fix this before coding if they show another person.",
        PALE_GREEN,
    )

    add_heading(doc, "3. Update main and create the feature branch", 2)
    add_command(
        doc,
        f"git checkout main\ngit pull origin main\ngit checkout -b {branch}\ngit branch --show-current",
    )
    add_callout(doc, "Expected result", f"The last line is {branch}.", PALE_GREEN)
    doc.add_paragraph(
        "If Git says the branch already exists, do not recreate it. Run:"
    )
    add_command(doc, f"git checkout {branch}")

    add_heading(doc, "4. Install dependencies and open VS Code", 2)
    add_command(doc, "npm install\ncode .")
    add_callout(
        doc,
        "PowerShell reminder",
        "Commands are typed only when the prompt is visible. Do not type output text, a branch name by itself, or the letter y after a command has already finished.",
        PALE_ORANGE,
    )


def add_validation_and_git(doc, branch, commit_title, commit_body, add_paths, pr_title, pr_body, tests):
    add_heading(doc, "Verify the feature before committing", 1)
    add_heading(doc, "1. Run the application", 2)
    add_command(doc, "npm run dev")
    doc.add_paragraph(
        "Open http://localhost:5173/ in Chrome. Work through every manual test below. Keep DevTools Console open and confirm there are no red errors."
    )
    add_checklist(doc, tests)
    doc.add_paragraph("Stop the development server with Ctrl+C, then run:")
    add_command(doc, "npm run build\ngit diff --check\ngit status --short")
    add_callout(
        doc,
        "Expected result",
        "The Vite build ends with 'built in ...'. git diff --check prints nothing. git status shows only the files from this feature.",
        PALE_GREEN,
    )

    add_heading(doc, "Commit and push using your own account", 1)
    doc.add_paragraph(
        "Review the changed-file list first. Do not use git add . for this assignment handoff; add only the intended files."
    )
    add_command(
        doc,
        f"git status --short\ngit add {add_paths}\n"
        f'git commit -m "{commit_title}" -m "{commit_body}"\n'
        f"git push -u origin {branch}\ngit status --short",
    )
    add_callout(
        doc,
        "Expected result",
        "Git reports a new commit, the branch is pushed to origin, and the final status is clean. Save a screenshot of this successful terminal output for the team report if requested.",
        PALE_GREEN,
    )

    add_heading(doc, "Open the pull request", 1)
    add_numbered(
        doc,
        [
            "Open the Huddle repository on GitHub.",
            "Click Compare & pull request for your branch.",
            "Confirm base is main and compare is your feature branch.",
            "Paste the title and description below.",
            "Request Dew3120 (the team leader) as reviewer.",
            "Create the pull request. Do not merge it yourself.",
            "Respond to any requested changes, push the fixes to the same branch, and tell the leader when it is ready again.",
        ],
    )
    add_code(doc, pr_title, "Pull-request title")
    add_code(doc, pr_body, "Pull-request description")
    add_callout(
        doc,
        "Merge rule",
        "The leader reviews the Files changed tab and merges with Create a merge commit. Do not squash, because the coursework asks for intact commit history.",
        PALE_ORANGE,
    )


def add_troubleshooting(doc):
    add_heading(doc, "Troubleshooting", 1)
    rows = [
        ("'npm' cannot find package.json", "You are in the wrong folder. Run cd C:\\Users\\<Windows-user>\\huddle, then retry."),
        ("The site cannot be reached", "Run npm run dev and leave that terminal running. Open the exact Local URL Vite prints."),
        ("The port is already in use", "Stop the old Vite terminal with Ctrl+C, or use the different Local URL Vite selects."),
        ("Branch already exists", "Run git checkout <branch-name>. Do not run checkout -b again."),
        ("Nothing to commit", "Save the files in VS Code, then run git status --short. Confirm you edited the correct repository."),
        ("LF will be replaced by CRLF", "This Windows line-ending warning is normal. It is not a failed commit."),
        ("A merge conflict appears", "Stop. Do not delete code. Send the conflict output to the team leader before continuing."),
        ("node_modules or dist appears in Git", "Do not add it. Confirm .gitignore includes node_modules/ and dist/."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(2.25)
    table.columns[1].width = Inches(4.75)
    headers = table.rows[0].cells
    headers[0].text = "Problem"
    headers[1].text = "What to do"
    set_repeat_table_header(table.rows[0])
    for cell in headers:
        shade_cell(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for problem, solution in rows:
        cells = table.add_row().cells
        cells[0].text = problem
        cells[1].text = solution
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_final_understanding(doc, explanation_points):
    add_heading(doc, "What you must be able to explain", 1)
    doc.add_paragraph(
        "Before the pull request is accepted, explain the points below to the leader in your own words. "
        "Do not memorize a sentence without understanding what the code does."
    )
    add_bullets(doc, explanation_points)
    add_heading(doc, "Personal completion checklist", 2)
    add_checklist(
        doc,
        [
            "I used my own Git name and verified GitHub email.",
            "I worked only on my assigned feature branch.",
            "I typed/read the code and can explain the feature.",
            "I manually tested every required behavior.",
            "npm run build passed.",
            "git diff --check printed no errors.",
            "I committed and pushed from my own account.",
            "I opened a pull request into main and requested the leader's review.",
            "I did not merge my own pull request.",
        ],
    )
    add_callout(
        doc,
        "Definition of done",
        "The feature works, the branch is pushed, the pull request is reviewed, and you can explain your contribution. A commit alone is not enough.",
        PALE_GREEN,
    )


def create_charles_guide():
    member = "Charles"
    branch = "feature/shared-button"
    doc = configure_document(f"{member} Guide")
    add_cover(
        doc,
        member,
        "Frontend contributor",
        branch,
        "Shared reusable Button component",
    )
    add_contents(
        doc,
        [
            ("Set up Git safely", "Use Charles's own identity, update main, and create the correct branch."),
            ("Build the component", "Create Button.jsx and Button.module.css with reusable variants."),
            ("Integrate it", "Replace repeated raw buttons without changing existing behavior."),
            ("Verify it", "Test every button, build the app, and check the diff."),
            ("Submit it", "Commit, push, open a PR, and request the leader's review."),
        ],
    )
    add_common_setup(
        doc,
        member,
        branch,
        "Begin only after the leader has merged the frontend-assignment branch into main. Tell the leader before starting so main is confirmed current.",
    )

    add_heading(doc, "Your assigned feature", 1)
    doc.add_paragraph(
        "The lecture requires one shared Button component with at least two variants and asks the team to use it wherever the app currently has a raw <button>. "
        "Your contribution satisfies that architecture requirement while keeping the existing Huddle behavior unchanged."
    )
    add_bullets(
        doc,
        [
            "Create one reusable Button component.",
            "Support primary, secondary, and danger appearances.",
            "Support medium and small sizes.",
            "Forward standard button properties such as type, onClick, disabled, and aria-label.",
            "Replace raw buttons in add, edit, move, delete, and retry actions.",
            "Preserve keyboard focus, hover, and disabled states.",
        ],
    )

    add_heading(doc, "Create the shared Button component", 1)
    add_heading(doc, "1. Create the folder and files", 2)
    add_command(
        doc,
        "New-Item -ItemType Directory -Force .\\src\\components\\Button\n"
        "code .\\src\\components\\Button\\Button.jsx\n"
        "code .\\src\\components\\Button\\Button.module.css",
    )
    add_heading(doc, "2. Paste Button.jsx", 2)
    add_code(
        doc,
        """import styles from './Button.module.css';

export default function Button({
  variant = 'primary',
  size = 'medium',
  className = '',
  children,
  ...props
}) {
  const classes = [
    styles.button,
    styles[variant],
    styles[size],
    className,
  ]
    .filter(Boolean)
    .join(' ');

  return (
    <button className={classes} {...props}>
      {children}
    </button>
  );
}
""",
        "src/components/Button/Button.jsx",
    )
    add_callout(
        doc,
        "Why this works",
        "variant and size choose CSS classes; children is the label; ...props forwards native button properties. This prevents every caller from needing a separate wrapper prop.",
        PALE_BLUE,
    )

    add_heading(doc, "3. Paste Button.module.css", 2)
    add_code(
        doc,
        """.button {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-height: 42px;
  padding: 9px 18px;
  border: 1px solid transparent;
  border-radius: 6px;
  font: inherit;
  font-weight: 700;
  line-height: 1;
  cursor: pointer;
  transition:
    background-color 150ms ease,
    border-color 150ms ease,
    color 150ms ease;
}

.button:focus-visible {
  outline: 3px solid rgb(53 120 168 / 25%);
  outline-offset: 2px;
}

.button:disabled {
  cursor: not-allowed;
  opacity: 0.45;
}

.primary {
  border-color: #176b55;
  background: #176b55;
  color: #ffffff;
}

.primary:hover:not(:disabled) {
  border-color: #125744;
  background: #125744;
}

.secondary {
  border-color: #b8c5d4;
  background: #ffffff;
  color: #1e3a5f;
}

.secondary:hover:not(:disabled) {
  border-color: #3578a8;
  background: #edf5fb;
}

.danger {
  border-color: #e7b6b6;
  background: #ffffff;
  color: #a51d2d;
}

.danger:hover:not(:disabled) {
  border-color: #d99090;
  background: #fff1f1;
}

.small {
  min-height: 36px;
  padding: 7px 10px;
  font-size: 0.8rem;
}

.medium {
  min-height: 42px;
  padding: 9px 18px;
}
""",
        "src/components/Button/Button.module.css",
    )

    add_heading(doc, "Integrate Button into the existing app", 1)
    doc.add_paragraph(
        "Open each file listed below. Add the import near the top, then replace only the shown raw button. Keep the surrounding form and handler code unchanged."
    )
    add_heading(doc, "1. AddTaskForm.jsx", 2)
    add_code(doc, "import Button from './Button/Button';", "Add with the other imports")
    add_code(doc, '<Button type="submit">Add task</Button>', "Replace the submit button")

    add_heading(doc, "2. EditTaskForm.jsx", 2)
    add_code(doc, "import Button from './Button/Button';", "Add with the other imports")
    add_code(
        doc,
        """<Button type="submit">Save changes</Button>
<Button type="button" variant="secondary" onClick={onCancel}>
  Cancel
</Button>""",
        "Replace the Save changes and Cancel buttons",
    )

    add_heading(doc, "3. TaskCard.jsx", 2)
    add_code(doc, "import Button from './Button/Button';", "Add with the other imports")
    add_code(
        doc,
        """<Button
  type="button"
  variant="secondary"
  size="small"
  onClick={() => onMove(task.id, previousStatus)}
  disabled={!previousStatus}
>
  Move left
</Button>""",
        "Move left button (retain the variable names already used in the file)",
    )
    add_code(
        doc,
        """<Button
  type="button"
  variant="secondary"
  size="small"
  onClick={() => onMove(task.id, nextStatus)}
  disabled={!nextStatus}
>
  Move right
</Button>""",
        "Move right button",
    )
    add_code(
        doc,
        """<Button
  type="button"
  variant="danger"
  size="small"
  className="task-card__delete"
  onClick={() => onDelete(task.id)}
>
  Delete
</Button>""",
        "Delete button",
    )
    add_callout(
        doc,
        "Do not guess handler names",
        "If the current TaskCard uses different callback or status variable names, preserve those exact expressions and change only the <button> tag to <Button> plus the variant/size props.",
        PALE_ORANGE,
    )

    add_heading(doc, "4. ErrorState.jsx", 2)
    add_code(doc, "import Button from './Button/Button';", "Add with the other imports")
    add_code(
        doc,
        """<Button type="button" variant="secondary" onClick={onRetry}>
  Try again
</Button>""",
        "Replace the retry button",
    )

    add_heading(doc, "Clean old global button styles", 1)
    doc.add_paragraph(
        "Open src/index.css. The Button module now owns appearance. Remove the old appearance blocks listed below so they do not fight the module styles:"
    )
    add_bullets(
        doc,
        [
            ".task-form button",
            ".task-form button:hover",
            ".task-card__actions button",
            ".task-card__actions button:hover:not(:disabled)",
            ".task-card__actions button:disabled",
            ".task-card__actions .task-card__delete appearance declarations",
            ".task-card__actions .task-card__delete:hover",
        ],
    )
    doc.add_paragraph("Keep this layout-only rule:")
    add_code(
        doc,
        """.task-card__actions .task-card__delete {
  margin-left: auto;
}""",
    )
    doc.add_paragraph("Then check that no raw button remains outside the shared component:")
    add_command(doc, 'rg -n "<button|</button>" .\\src')
    add_callout(
        doc,
        "Expected result",
        "Only src/components/Button/Button.jsx contains a raw <button> element.",
        PALE_GREEN,
    )

    add_validation_and_git(
        doc,
        branch,
        "Add reusable button component",
        "Create a shared Button with primary, secondary, and danger variants, replace repeated raw buttons, and keep accessible focus and disabled states consistent.",
        "src/components/Button src/components/AddTaskForm.jsx src/components/EditTaskForm.jsx src/components/TaskCard.jsx src/components/ErrorState.jsx src/index.css",
        "Add shared reusable Button component",
        """## Summary
- added a reusable Button component with primary, secondary, and danger variants
- replaced repeated raw buttons in task forms, task cards, and retry state
- kept focus, hover, disabled, and small-size behavior consistent

## Verification
- `npm run build`
- `git diff --check`
- manually tested add, edit, cancel, move, delete, retry, and disabled buttons""",
        [
            "Add a valid task and confirm the primary Add task button works.",
            "Open a task, edit it, and confirm Save changes works.",
            "Open edit mode again and confirm Cancel returns without saving.",
            "Move tasks left and right; boundary movement buttons must be visibly disabled.",
            "Delete a task and confirm the delete action still works.",
            "Confirm the danger button is red, secondary buttons are outlined, and the primary button is green.",
            "Use the Tab key and confirm each button has a visible focus outline.",
            "Confirm the browser console has no red errors or React warnings.",
        ],
    )
    add_troubleshooting(doc)
    add_final_understanding(
        doc,
        [
            "A shared Button removes duplicated markup and styles and gives the app consistent behavior.",
            "variant changes meaning/appearance (primary, secondary, danger); size changes dimensions.",
            "children is the visible button content.",
            "...props forwards native properties such as type, onClick, disabled, and accessibility attributes.",
            "CSS Modules scope the Button styles so they do not accidentally style unrelated elements.",
            "The component still renders a real HTML button, so keyboard and form behavior remain correct.",
        ],
    )
    path = OUTPUT_DIR / "Charles-Assignment-01-Guide.docx"
    doc.save(path)
    return path


def create_vinuka_guide():
    member = "Vinuka"
    branch = "feature/task-filters"
    doc = configure_document(f"{member} Guide")
    add_cover(
        doc,
        member,
        "Frontend contributor",
        branch,
        "Task search and URL-backed filters",
    )
    add_contents(
        doc,
        [
            ("Set up Git safely", "Use Vinuka's own identity and branch from the updated main."),
            ("Build filter utilities", "Add pure, testable assignee and overdue filtering logic."),
            ("Build filter controls", "Create title, assignee, status, and overdue controls."),
            ("Integrate URL state", "Keep filters after refresh and make filtered views shareable."),
            ("Verify and submit", "Test combinations, build, push, and open a reviewed PR."),
        ],
    )
    add_common_setup(
        doc,
        member,
        branch,
        "Begin only after the leader has merged Charles's shared Button pull request into main. Your TaskFilters component imports that shared Button.",
    )

    add_heading(doc, "Your assigned feature", 1)
    doc.add_paragraph(
        "The Assignment 1 slides require filtering by assignee or status, title search, a proper no-results state, and URL-backed filter state. "
        "Your contribution adds all of these without changing task creation, movement, editing, or deletion."
    )
    add_bullets(
        doc,
        [
            "Search task titles without case sensitivity.",
            "Filter by assignee.",
            "Filter by status.",
            "Show only overdue, unfinished tasks.",
            "Combine multiple filters.",
            "Reflect filter state in the URL query string.",
            "Restore the same filters after a browser refresh.",
            "Show a clear no-results message and a Clear filters action.",
        ],
    )

    add_heading(doc, "Create pure filter utilities", 1)
    add_heading(doc, "1. Open the utility file", 2)
    add_command(doc, "code .\\src\\utils\\filterTasks.js")
    add_heading(doc, "2. Paste the complete utility code", 2)
    add_code(
        doc,
        """export function getAssignees(tasks) {
  return [...new Set(tasks.map((task) => task.assignee).filter(Boolean))].sort(
    (a, b) => a.localeCompare(b),
  );
}

export function isOverdue(task, today = new Date()) {
  if (!task.dueDate || task.status === 'done') {
    return false;
  }

  const dueDate = new Date(`${task.dueDate}T23:59:59`);
  return dueDate < today;
}

export function filterTasks(tasks, filters) {
  const query = filters.query.trim().toLowerCase();

  return tasks.filter((task) => {
    const matchesQuery =
      query.length === 0 || task.title.toLowerCase().includes(query);
    const matchesAssignee =
      filters.assignee.length === 0 || task.assignee === filters.assignee;
    const matchesStatus =
      filters.status.length === 0 || task.status === filters.status;
    const matchesOverdue = !filters.overdue || isOverdue(task);

    return (
      matchesQuery &&
      matchesAssignee &&
      matchesStatus &&
      matchesOverdue
    );
  });
}
""",
        "src/utils/filterTasks.js",
    )
    add_callout(
        doc,
        "Why this is separate",
        "These are pure data functions, so the UI stays readable and the logic will be easy to unit test in Session 4. Done tasks are not overdue even if their date is old.",
        PALE_BLUE,
    )

    add_heading(doc, "Create the TaskFilters component", 1)
    add_heading(doc, "1. Open the component file", 2)
    add_command(doc, "code .\\src\\components\\TaskFilters.jsx")
    add_heading(doc, "2. Paste the complete component", 2)
    add_code(
        doc,
        """import Button from './Button/Button';

export default function TaskFilters({
  filters,
  assignees,
  onChange,
  onClear,
}) {
  return (
    <section className="task-filters" aria-label="Task filters">
      <label>
        Search
        <input
          type="search"
          value={filters.query}
          onChange={(event) => onChange('query', event.target.value)}
          placeholder="Search task titles"
        />
      </label>

      <label>
        Assignee
        <select
          value={filters.assignee}
          onChange={(event) => onChange('assignee', event.target.value)}
        >
          <option value="">All assignees</option>
          {assignees.map((assignee) => (
            <option key={assignee} value={assignee}>
              {assignee}
            </option>
          ))}
        </select>
      </label>

      <label>
        Status
        <select
          value={filters.status}
          onChange={(event) => onChange('status', event.target.value)}
        >
          <option value="">All statuses</option>
          <option value="todo">To Do</option>
          <option value="in-progress">In Progress</option>
          <option value="done">Done</option>
        </select>
      </label>

      <label className="task-filters__checkbox">
        <input
          type="checkbox"
          checked={filters.overdue}
          onChange={(event) => onChange('overdue', event.target.checked)}
        />
        Overdue only
      </label>

      <Button type="button" variant="secondary" onClick={onClear}>
        Clear filters
      </Button>
    </section>
  );
}
""",
        "src/components/TaskFilters.jsx",
    )

    add_heading(doc, "Integrate the filters into Board.jsx", 1)
    doc.add_paragraph(
        "Open src/components/Board.jsx. Keep all existing task handlers and the done counter. Apply the following changes carefully."
    )
    add_heading(doc, "1. Add imports", 2)
    add_code(
        doc,
        """import { useSearchParams } from 'react-router-dom';
import TaskFilters from './TaskFilters';
import { filterTasks, getAssignees } from '../utils/filterTasks';""",
    )

    add_heading(doc, "2. Read filters from the URL", 2)
    doc.add_paragraph("Inside the Board function, after the existing task state/context values, add:")
    add_code(
        doc,
        """const [searchParams, setSearchParams] = useSearchParams();

const filters = {
  query: searchParams.get('q') ?? '',
  assignee: searchParams.get('assignee') ?? '',
  status: searchParams.get('status') ?? '',
  overdue: searchParams.get('overdue') === 'true',
};

const assignees = getAssignees(tasks);
const visibleTasks = filterTasks(tasks, filters);""",
    )

    add_heading(doc, "3. Add URL update functions", 2)
    add_code(
        doc,
        """function updateFilters(name, value) {
  const nextParams = new URLSearchParams(searchParams);
  const key = name === 'query' ? 'q' : name;

  if (value === '' || value === false) {
    nextParams.delete(key);
  } else {
    nextParams.set(key, String(value));
  }

  setSearchParams(nextParams, { replace: true });
}

function clearFilters() {
  setSearchParams({}, { replace: true });
}""",
    )

    add_heading(doc, "4. Render the filter controls", 2)
    doc.add_paragraph("Place this above the board columns:")
    add_code(
        doc,
        """<TaskFilters
  filters={filters}
  assignees={assignees}
  onChange={updateFilters}
  onClear={clearFilters}
/>""",
    )

    add_heading(doc, "5. Render filtered tasks and no-results state", 2)
    doc.add_paragraph(
        "Where the existing code filters tasks for each column, use visibleTasks instead of tasks. Wrap the existing board markup with this condition:"
    )
    add_code(
        doc,
        """{visibleTasks.length === 0 ? (
  <section className="empty-state" role="status">
    <h2>No matching tasks</h2>
    <p>Change or clear the filters to see tasks again.</p>
  </section>
) : (
  <div className="board">
    {/* Keep the existing mapped columns here.
        Each column must filter visibleTasks, not tasks. */}
  </div>
)}""",
    )
    add_callout(
        doc,
        "Do not delete existing handlers",
        "The Board must still add, move, edit, and delete tasks. Change only the array used to display columns. The done counter should continue to use the full tasks array, not visibleTasks.",
        PALE_ORANGE,
    )

    add_heading(doc, "Style the filters and empty state", 1)
    doc.add_paragraph("Append this to src/index.css:")
    add_code(
        doc,
        """.task-filters {
  display: grid;
  grid-template-columns: 2fr 1fr 1fr auto auto;
  align-items: end;
  gap: 14px;
  max-width: 1400px;
  margin: 0 auto 24px;
  padding: 18px;
  border: 1px solid #d7e0ea;
  border-radius: 8px;
  background: #ffffff;
}

.task-filters label {
  display: grid;
  gap: 7px;
  color: #334155;
  font-size: 0.875rem;
  font-weight: 600;
}

.task-filters input,
.task-filters select {
  width: 100%;
  min-height: 42px;
  padding: 9px 11px;
  border: 1px solid #cbd5e1;
  border-radius: 6px;
  background: #ffffff;
  color: #0f172a;
  font: inherit;
}

.task-filters__checkbox {
  display: flex !important;
  align-items: center;
  min-height: 42px;
  white-space: nowrap;
}

.task-filters__checkbox input {
  width: 18px;
  min-height: 18px;
  margin: 0;
}

.empty-state {
  max-width: 1400px;
  margin: 0 auto;
  padding: 32px;
  border: 1px dashed #b8c5d4;
  border-radius: 8px;
  background: #ffffff;
  text-align: center;
}

.empty-state h2 {
  margin: 0 0 8px;
}

.empty-state p {
  margin: 0;
  color: #667085;
}

@media (max-width: 1000px) {
  .task-filters {
    grid-template-columns: 1fr 1fr;
  }
}

@media (max-width: 600px) {
  .task-filters {
    grid-template-columns: 1fr;
  }
}
""",
        "Append to src/index.css",
    )

    add_validation_and_git(
        doc,
        branch,
        "Add task search and filters",
        "Add title search and assignee, status, and overdue filters, store filter state in the URL, and show a clear no-results state.",
        "src/components/TaskFilters.jsx src/components/Board.jsx src/utils/filterTasks.js src/index.css",
        "Add task search and URL-backed filters",
        """## Summary
- added title search and assignee, status, and overdue filters
- stored filter state in the URL query string
- added clear-filter and no-results behavior
- kept existing task actions working

## Verification
- `npm run build`
- `git diff --check`
- manually tested individual and combined filters, refresh persistence, clear filters, and no-results state""",
        [
            "Type part of a title in Search; matching is case-insensitive.",
            "Choose each assignee and confirm only that person's tasks appear.",
            "Choose To Do, In Progress, and Done and confirm each status filter works.",
            "Enable Overdue only; old unfinished tasks appear, while completed tasks do not.",
            "Combine title, assignee, and status filters and confirm all conditions apply together.",
            "Confirm the browser URL changes (for example ?q=login&status=todo).",
            "Refresh the browser and confirm the same filters remain selected.",
            "Choose filters that match nothing and confirm No matching tasks appears.",
            "Click Clear filters and confirm all tasks and columns return.",
            "After filtering, confirm add, edit, move, and delete still work.",
            "Resize to a narrow window and confirm the filter controls stack without overlapping.",
            "Confirm the browser console has no red errors or React warnings.",
        ],
    )
    add_troubleshooting(doc)
    add_final_understanding(
        doc,
        [
            "Filters are derived from the task array; they do not modify or duplicate task state.",
            "URLSearchParams stores filter choices in the query string, which makes refresh and sharing work.",
            "filterTasks combines every active condition with logical AND.",
            "getAssignees uses a Set to build a sorted list without duplicate names.",
            "isOverdue ignores completed tasks and compares unfinished due dates with today.",
            "The no-results state means data loaded successfully but the current filters match nothing.",
            "The done counter uses all tasks, while visibleTasks controls only what the board displays.",
        ],
    )
    path = OUTPUT_DIR / "Vinuka-Assignment-01-Guide.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_charles_guide()
    create_vinuka_guide()
