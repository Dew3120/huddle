from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
IMAGE_DIR = REPO / "docs" / "screenshots" / "assignment-03"
OUTPUT = DESKTOP / "Group-97-Huddle-Assignment-03-Full-Stack-Application-Report.docx"

NAVY = "123047"
BLUE = "4560B2"
ORANGE = "E58A2B"
LIGHT_BLUE = "EAF0F7"
LIGHT_ORANGE = "FFF4E7"
MID_GREY = "6B7280"
DARK = "263443"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="CBD5E1", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, value, size=8.5, bold=False, color=DARK):
    cell.text = ""
    lines = str(value).split("\n")
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        run.font.name = "Aptos"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths=None, size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, heading in enumerate(headers):
        set_cell_text(header.cells[index], heading, size=size, bold=True, color="FFFFFF")
        set_cell_shading(header.cells[index], NAVY)
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for index, value in enumerate(row_data):
            set_cell_text(row.cells[index], value, size=size)
            set_cell_shading(row.cells[index], "F8FAFC" if row_index % 2 else "FFFFFF")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(result)
    run._r.append(end)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    relationship_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_properties.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF3F7")
    p_pr.append(shd)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(DARK)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)
    return paragraph


def add_note(doc, label, text, fill=LIGHT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, ORANGE, "10")
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    first = paragraph.add_run(f"{label}: ")
    first.bold = True
    first.font.color.rgb = RGBColor.from_string(ORANGE)
    paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_figure(doc, number, filename, caption):
    path = IMAGE_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(8)
    label = heading.add_run(f"Figure {number}")
    label.bold = True
    label.font.size = Pt(11)
    label.font.color.rgb = RGBColor.from_string(ORANGE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    with Image.open(path) as image:
        image_width, image_height = image.size
    ratio = image_width / image_height
    width = min(6.45, 5.65 * ratio)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor.from_string(MID_GREY)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (("Heading 1", 17, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11, ORANGE)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name != "Heading 1" else 14)
        style.paragraph_format.space_after = Pt(5)

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title_p_pr = title.element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    header = section.header.paragraphs[0]
    header.text = "GROUP 97  |  HUDDLE / SYNCBOARD"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MID_GREY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = footer.add_run("Assignment 03  |  ")
    left.font.name = "Aptos"
    left.font.size = Pt(8)
    left.font.color.rgb = RGBColor.from_string(MID_GREY)
    add_field(footer, "PAGE")
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GREY)


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(42)
    run = paragraph.add_run("ASSIGNMENT 03")
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ORANGE)
    run.font.letter_spacing = None

    title = doc.add_paragraph(style="Title")
    title.add_run("Working Full Stack Application")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(2)
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("Frontend, Backend and Database")
    run.font.name = "Aptos Display"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(18)
    run = line.add_run("HUDDLE / SYNCBOARD CLIENT PROJECT")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    add_note(doc, "Group", "97", fill=LIGHT_BLUE)
    add_note(doc, "Prepared", "6 September 2026", fill="F8FAFC")
    add_note(doc, "Submission state", "Full-stack implementation verified against Atlas Free. Final release evidence is included.", fill=LIGHT_BLUE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(70)
    paragraph = doc.add_paragraph()
    paragraph.add_run("Repository: ")
    add_hyperlink(paragraph, "github.com/Dew3120/huddle", "https://github.com/Dew3120/huddle")
    paragraph = doc.add_paragraph()
    paragraph.add_run("Release tag: ")
    paragraph.add_run("assignment-03-working-full-stack-application").bold = True

    doc.add_page_break()


def add_toc(doc):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Title"]
    paragraph.add_run("Table of Contents")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph("The table of contents is a Word field. Open the DOCX in Word and update the field if pagination changes.")
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Group 97 - Assignment 03 Working Full Stack Application"
    doc.core_properties.subject = "Huddle / SyncBoard frontend, backend and database report"
    doc.core_properties.author = "Group 97"
    doc.core_properties.keywords = "Assignment 03, Huddle, SyncBoard, MongoDB, Mongoose, PouchDB"

    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "Document Control", 1)
    add_table(doc, ["Item", "Value"], [
        ("Module deliverable", "Assignment 03 - Working Full Stack Application (Frontend, Backend and Database)"),
        ("Project", "Huddle / SyncBoard collaborative task board"),
        ("Group", "97"),
        ("Repository", "https://github.com/Dew3120/huddle"),
        ("Working integration branch", "feature/session-03-client-persistence"),
        ("Final release tag", "assignment-03-working-full-stack-application"),
        ("Evidence date", "6 September 2026"),
    ], widths=[1.7, 4.9], size=9)

    add_heading(doc, "1. Introduction", 1)
    doc.add_paragraph(
        "Huddle, presented as the SyncBoard client project, is a collaborative task board for creating, assigning, moving, filtering, and tracking team work. The Assignment 03 implementation connects the React frontend to a Node.js and Express backend backed by MongoDB through Mongoose. It also adds browser-side persistence with PouchDB so the board can render from local data, continue during a temporary API outage, and replay queued changes when connectivity returns."
    )
    doc.add_paragraph(
        "The implementation keeps the application observable for a student demonstration: the API exposes health and documentation endpoints, the data model has explicit schemas and indexes, the verification scripts exercise the main HTTP contract, and the client presents pending, failed, and conflicting task states instead of silently discarding edits."
    )
    add_note(doc, "Scope note", "The implementation was verified end to end against the team's Atlas Free deployment. Credentials and the private connection URI are intentionally excluded from this report.")

    add_heading(doc, "2. Requirements Traceability", 1)
    add_table(doc, ["Assignment requirement", "Evidence in this submission", "Status"], [
        ("Frontend", "React routes, board, task creation/editing, filters, authentication and offline UI; Figures 1-16.", "Complete"),
        ("Backend", "Express routes, controllers, services, repositories, validation, Swagger and health endpoint; Figures 17, 18, 20, 21 and 24.", "Complete"),
        ("Database", "Mongoose models, Atlas persistence, indexes, aggregation and database documents; Figures 19, 22 and 23.", "Complete"),
        ("Introduction and team roles", "Sections 1 and 3, including NSBM IDs and contribution evidence.", "Complete"),
        ("Atlas Free database account", "Section 11 and Figures 22-24 show the Free cluster, seeded collections and connected API.", "Complete"),
        ("GitHub link and tagged release", "Section 12 identifies the repository, contributor history and annotated release tag.", "Complete"),
        ("README how to run", "The repository README contains local, Atlas, seed, server, client, test and offline instructions.", "Complete"),
        ("Postman or Swagger API collection", "Committed Postman collection, OpenAPI YAML, Swagger UI and evidence screenshots.", "Complete"),
        ("Screenshots of frontend, backend and database", "Figures 1-24 cover all application pages, API behavior, local persistence and Atlas evidence.", "Complete"),
    ], widths=[2.0, 3.8, 0.8], size=8.2)

    add_heading(doc, "3. Team and Contributions", 1)
    doc.add_paragraph("The names below use the NSBM ID in brackets whenever a team member is mentioned. Plymouth IDs are included where they were provided to the group.")
    add_table(doc, ["Member", "NSBM ID", "Plymouth ID", "Assignment 03 role"], [
        ("T D Gnanasena (NSBM 36407)", "36407", "10967149", "Project coordination, MongoDB/Mongoose integration, authentication and CRUD persistence, integration of client-persistence work, PWA/offline completion, README, report and release preparation."),
        ("J Charles (NSBM 36359)", "36359", "10967132", "Assignment 03 verification contribution, including the repeatable persistence/conflict verification script and clean test/build/audit evidence."),
        ("K V Dilnath (NSBM 33700)", "33700", "10965490", "PouchDB browser cache and user-scoped task persistence, including cached-first loading and the client-persistence branch contribution."),
        ("R S Bokalagama (NSBM 37412)", "37412", "Not provided", "Task statistics aggregation and optimistic task-concurrency behavior, including stale-version conflict evidence."),
    ], widths=[1.75, 0.75, 0.9, 3.2], size=8.2)

    add_heading(doc, "3.1 Commit evidence", 2)
    doc.add_paragraph("The following named commits identify the main Assignment 03 contributions in the Git history. Full hashes remain available in the repository log.")
    add_table(doc, ["Member", "Named commits and contribution evidence"], [
        ("T D Gnanasena (NSBM 36407)", "Connect the API to MongoDB (cb4462a); Define user and board persistence models (3897287); Define task schema and query indexes (d55eb29); Add repeatable MongoDB demo seeding (828d5cd); Persist authentication users with Mongoose (495146d); Persist boards with Mongoose (b1f358d); Persist task queries and CRUD with Mongoose (ffb9bc3); Map duplicate user emails to conflict responses (8859581); Reject malformed MongoDB resource IDs (c699733); Correct the task collection API assertion (3f1de13); Serialize public users through the model (5f68bc2); Update the qs security patch (642fe8a); Merge Vinuka PouchDB task cache (5ba2816); Queue offline task changes for synchronization (23dff0d); Show offline and task conflict states (5a1652c); Document offline synchronization workflow (ae0738b); Add a production offline app shell (a1695b1); Complete offline task synchronization (39705c9); Document Assignment 03 evidence (29883fc); Make Atlas connectivity resilient (589d8f0)."),
        ("J Charles (NSBM 36359)", "Add Assignment 03 persistence verification (c8e5f98), on feature/charles-assignment-03-verification. The contribution verifies MongoDB connection, persisted boards/tasks, version increments, stale-write rejection, aggregation results and cleanup."),
        ("K V Dilnath (NSBM 33700)", "Add a user-scoped PouchDB task cache (4bfdf60); Load tasks from PouchDB before refreshing (0f3a7e9). These commits provide the browser-local task data foundation."),
        ("R S Bokalagama (NSBM 37412)", "Reject stale task updates with version checks (c391b7f); Add board task statistics aggregation (faea66c). These commits provide the server-side concurrency and aggregation behavior used by the verification evidence."),
    ], widths=[1.75, 4.9], size=7.8)

    add_heading(doc, "4. System Architecture", 1)
    doc.add_paragraph("The application follows a layered client/server structure. The controller is the HTTP boundary; services contain business rules and ownership checks; repositories own MongoDB access; and the Mongoose models define the persisted shape.")
    add_code(doc, "React pages and providers\n        |\n        +-- API client with JWT and error handling\n        |\n        +-- Express routes -> controllers -> services -> repositories\n        |                                      |\n        |                                      +-- Mongoose models\n        |                                              |\n        |                                              +-- MongoDB\n        |\n        +-- PouchDB over IndexedDB for cached tasks and queued mutations")
    add_table(doc, ["Layer", "Responsibility", "Key files"], [
        ("Frontend", "Authentication, board rendering, task CRUD, filters and synchronization states.", "src/App.jsx, src/context, src/components, src/pages"),
        ("API client", "Bearer token, JSON requests, API errors, network classification and 204 handling.", "src/api/client.js, src/api/auth.js, src/api/tasks.js"),
        ("Backend", "Routes, request validation, authentication, ownership and response shaping.", "server/routes, server/controllers, server/services"),
        ("Persistence", "MongoDB connection, Mongoose schemas, indexes, repositories and seed data.", "server/db, server/models, server/repositories"),
        ("Browser persistence", "User-scoped cache, queued changes, reconnect replay and conflict state.", "src/db/taskCache.js, src/services/taskSynchronization.js"),
    ], widths=[1.1, 3.2, 2.3], size=8.4)

    add_heading(doc, "5. Database Design and Persistence", 1)
    doc.add_paragraph("Tasks remain separate documents because they are updated frequently, queried independently, and can grow without a fixed bound. Boards embed their small, bounded column list. Activity is kept separate because it grows over time and should not enlarge the board document on every event.")
    add_table(doc, ["Collection", "Decision", "Reason"], [
        ("users", "Separate documents; referenced by boards/tasks.", "Authentication and independent user lookup."),
        ("boards", "Separate documents; columns embedded.", "Board ownership is an access boundary and columns are small and bounded."),
        ("tasks", "Separate documents with boardId and optional columnId.", "Independent filtering, sorting, pagination, updates and version checks."),
        ("activity", "Separate documents with boardId, taskId and userId.", "Activity grows continuously and should not enlarge core documents."),
    ], widths=[1.0, 2.5, 3.1], size=8.5)
    add_heading(doc, "5.1 Mongoose model decisions", 2)
    add_bullet(doc, "User emails are required, trimmed, lowercased and unique. The model's toJSON transform removes passwordHash and exposes a public id.")
    add_bullet(doc, "Task status and priority use enums; title and dueDate are validated; timestamps are enabled; version is an explicit integer used for optimistic concurrency.")
    add_bullet(doc, "Board columns and members use embedded subdocuments, while task board/user relationships use ObjectId references.")
    add_heading(doc, "5.2 Indexes and aggregation", 2)
    add_table(doc, ["Index or query", "Purpose"], [
        ("tasks_by_board_status_position", "Board screen equality filters followed by position ordering."),
        ("tasks_by_board_due_date", "Overdue and due-date queries by board."),
        ("tasks_by_assignee_status", "My-task and status views."),
        ("tasks_text_search", "Text search over title and description."),
        ("users.email unique", "Database-level one-account-per-email constraint."),
        ("GET /api/boards/:id/task-stats", "MongoDB aggregation with byStatus and overdueByAssignee facet groups."),
    ], widths=[2.5, 4.1], size=8.5)

    add_heading(doc, "6. Client Persistence and Conflict Handling", 1)
    doc.add_paragraph("Each signed-in user receives a separate PouchDB database in the browser. Server tasks are stored as task documents, while create/update/delete operations are stored as mutation documents. This separation keeps PouchDB revision metadata out of the task shape used by React.")
    add_bullet(doc, "The board reads cached tasks first and refreshes from the API when available.")
    add_bullet(doc, "Offline task writes remain visible immediately and are marked as pending instead of being lost.")
    add_bullet(doc, "Queued updates are compacted for the same task, and a queued create is updated in place rather than creating unnecessary requests.")
    add_bullet(doc, "When the API returns, queued mutations replay and the cache is refreshed. Gateway failures, offline status and network errors are treated as retryable conditions.")
    add_bullet(doc, "The server compares the submitted task version atomically. A stale write returns 409 TASK_CONFLICT with the current task and the version used by the client.")
    add_bullet(doc, "Non-overlapping field changes can be retried against the current version; same-field conflicts remain visible and require a user decision.")
    add_note(doc, "Conflict policy", "Huddle uses optimistic concurrency because it supports offline work without locking a task for every editor. The losing edit is reported to the user rather than silently overwriting another change.", fill=LIGHT_BLUE)

    add_heading(doc, "7. API Collection and Documentation", 1)
    doc.add_paragraph("The repository contains both a committed Postman collection and an OpenAPI/Swagger reference. The collection is the primary manual API evidence set; Swagger provides an interactive view of the same HTTP contract.")
    add_table(doc, ["Resource", "Location or endpoint"], [
        ("Postman collection", "docs/api-evidence/Huddle Assignment 02 API Evidence.postman_collection.json"),
        ("OpenAPI source", "docs/openapi.yaml"),
        ("Swagger UI", "http://localhost:4000/api/docs/"),
        ("OpenAPI JSON", "http://localhost:4000/api/openapi.json"),
        ("Health", "GET http://localhost:4000/api/health"),
        ("Task statistics", "GET /api/boards/:id/task-stats"),
    ], widths=[1.7, 4.9], size=8.5)
    add_heading(doc, "7.1 Main API operations", 2)
    add_table(doc, ["Area", "Operations demonstrated"], [
        ("Authentication", "Register, duplicate registration conflict, login and current-user lookup."),
        ("Boards", "List owned boards, create a board, list board tasks and enforce ownership."),
        ("Tasks", "List with filter/sort/pagination, create, read, update, delete and invalid-id handling."),
        ("Concurrency", "Versioned PATCH and 409 TASK_CONFLICT for stale writes."),
        ("Aggregation", "Status counts and overdue counts grouped by assignee."),
    ], widths=[1.4, 5.2], size=8.5)

    add_heading(doc, "8. How to Run", 1)
    doc.add_paragraph("The following commands are also in the repository README. Run the backend and frontend in separate PowerShell windows.")
    add_heading(doc, "8.1 Install and configure", 2)
    add_code(doc, "git clone https://github.com/Dew3120/huddle.git\ncd huddle\ngit switch feature/session-03-client-persistence\nnpm ci\nCopy-Item .env.example .env")
    add_code(doc, "PORT=4000\nJWT_SECRET=replace-this-with-a-long-random-development-secret\nCLIENT_ORIGIN=http://localhost:5173\nMONGODB_URI=mongodb://127.0.0.1:27017/huddle")
    add_heading(doc, "8.2 Seed and start", 2)
    add_code(doc, "npm run seed\n\n# PowerShell window 1\nnpm run dev:server\n\n# PowerShell window 2\nnpm run dev")
    doc.add_paragraph("Open the Vite URL, normally http://localhost:5173. Demo login: user1@nsbm.lk / password123.")
    add_heading(doc, "8.3 Verify and demonstrate", 2)
    add_code(doc, "npm test\nnpm run build\nnpm audit\nnpm run verify:assignment-02\nnpm run verify:assignment-03")
    add_bullet(doc, "For offline evidence, run the production preview after npm run build, load the board, stop the API, edit a task, refresh, restart the API, and choose Try reconnecting.")
    add_bullet(doc, "For Atlas, place the private Atlas URI only in the local .env, run the seed command, and confirm /api/health before taking evidence screenshots.")

    add_heading(doc, "9. Verification Results", 1)
    add_table(doc, ["Check", "Result", "Evidence"], [
        ("Client persistence tests", "10/10 passed", "npm test"),
        ("Assignment 02 API verification", "17/17 passed", "npm run verify:assignment-02"),
        ("Assignment 03 persistence verification", "9/9 passed", "npm run verify:assignment-03"),
        ("Production build", "Succeeded; PWA service-worker assets generated", "npm run build"),
        ("Dependency audit", "0 vulnerabilities", "npm audit --omit=optional"),
        ("Whitespace check", "No whitespace errors", "git diff --check"),
        ("Manual offline/reconnect/conflict flow", "Observed and captured", "Figures 11-16"),
        ("MongoDB health", "status ok, database connected, readyState 1", "GET /api/health and Figure 21"),
    ], widths=[2.1, 2.2, 2.3], size=8.5)
    add_note(doc, "Verification boundary", "The 17/17 API checks and 9/9 Assignment 03 checks were repeated against Atlas Free. The private Atlas URI remained only in the ignored local .env file.")

    add_heading(doc, "10. Screenshot Evidence", 1)
    doc.add_paragraph("The following figures are embedded from the repository's Assignment 03 evidence folder. Figures 1-10 cover the main frontend screens and task states. Figures 11-16 demonstrate client persistence, reconnection and conflict handling. Figures 17-21 cover backend and local database evidence. Figures 22-24 demonstrate the Atlas Free deployment, seeded collections and connected application health.")
    figure_specs = [
        ("01-sign-in.png", "Frontend sign-in screen."),
        ("02-sign-up.png", "Frontend sign-up screen."),
        ("03-board.png", "Authenticated task board loaded from the API."),
        ("04-create-task.png", "Create-task workflow."),
        ("05-task-detail.png", "Task detail view."),
        ("06-edit-task.png", "Edit-task workflow."),
        ("07-filtered-board.png", "Filtered board view."),
        ("08-no-results.png", "Empty filtered result state."),
        ("09-not-found.png", "Client not-found state."),
        ("10-delete-confirmation.png", "Delete confirmation state."),
        ("11-offline-pending.png", "Offline task edit remains visible and waits to synchronize."),
        ("12-offline-board.png", "Board remains usable while the API is unavailable."),
        ("13-offline-reload.png", "Reloaded app shell shows one queued change waiting to sync."),
        ("14-reconnected.png", "Reconnected board after queued synchronization."),
        ("15-conflict.png", "Task detail shows a conflicting edit state."),
        ("16-conflict-board.png", "Conflict resolution controls show the server and client versions."),
        ("17-backend-task-stats.png", "Backend aggregation response grouped by task status and overdue assignee."),
        ("18-backend-conflict-response.png", "Backend 409 TASK_CONFLICT response from the optimistic concurrency test."),
        ("19-mongodb-compass-tasks.png", "MongoDB Compass view of persisted task documents."),
        ("20-swagger-api-reference.png", "Swagger UI API reference."),
        ("21-backend-health-response.png", "Backend health response showing the connected database state."),
    ]
    for number, (filename, caption) in enumerate(figure_specs, start=1):
        add_figure(doc, number, filename, caption)
    add_figure(doc, 22, "22-atlas-free-cluster.png", "MongoDB Atlas Cluster0 overview showing the Free deployment in the Mumbai region.")
    add_figure(doc, 23, "23-atlas-huddle-collections.png", "Atlas Data Explorer showing the seeded huddle database with boards, tasks and users collections.")
    add_figure(doc, 24, "24-atlas-connected-health.png", "Application health endpoint confirming MongoDB is connected to Atlas with readyState 1.")

    add_heading(doc, "11. Atlas Free Database Account", 1)
    add_note(doc, "Current status", "Complete. Cluster0 is an Atlas Free deployment in AWS Mumbai, and the application was verified against its huddle database.")
    doc.add_paragraph("The Atlas application user is restricted to readWrite access on the huddle database. The developer IP address was added to the Atlas access list, the database was seeded through the project's repeatable seed command, and the API reported a connected Mongoose state. The password and connection URI remain private in the ignored local .env file.")
    add_table(doc, ["Atlas control", "Verified evidence"], [
        ("Account/project", "Group Atlas project and Cluster0 overview shown in Figure 22."),
        ("Free deployment", "Cluster0 is visibly labelled FREE and hosted in AWS Mumbai."),
        ("Least-privilege database user", "Application user is limited to readWrite on huddle; no credential is disclosed."),
        ("Seeded collections", "Atlas Data Explorer shows boards, tasks and users in Figure 23."),
        ("Connected application", "GET /api/health reports connected and readyState 1 in Figure 24."),
    ], widths=[2.2, 4.4], size=8.5)

    add_heading(doc, "12. GitHub and Release Tag", 1)
    paragraph = doc.add_paragraph()
    paragraph.add_run("Repository: ")
    add_hyperlink(paragraph, "https://github.com/Dew3120/huddle", "https://github.com/Dew3120/huddle")
    add_bullet(doc, "The working integration branch is feature/session-03-client-persistence.")
    add_bullet(doc, "The contributor commits for T D Gnanasena (NSBM 36407), J Charles (NSBM 36359), K V Dilnath (NSBM 33700), and R S Bokalagama (NSBM 37412) are visible in the Git history described in Section 3.1.")
    add_bullet(doc, "The reviewed submission is merged into main and identified by the annotated tag assignment-03-working-full-stack-application.")
    add_code(doc, "git fetch --tags\ngit checkout assignment-03-working-full-stack-application\n# Inspect this immutable Assignment 03 release\ngit checkout main")
    add_note(doc, "Release integrity", "The tag identifies the reviewed main-branch snapshot after Atlas verification, report completion and final automated checks.", fill=LIGHT_BLUE)

    add_heading(doc, "13. Conclusion and Submission Checklist", 1)
    doc.add_paragraph("The Huddle / SyncBoard client demonstrates the required full-stack behavior: the React frontend talks to an Atlas-backed MongoDB API, the API exposes documented CRUD and aggregation operations, and the client maintains a PouchDB task cache with queued offline edits and visible optimistic-concurrency conflicts.")
    add_table(doc, ["Final action", "State"], [
        ("Frontend, backend and Atlas database implementation", "Complete and verified"),
        ("README run instructions", "Updated for local, Atlas, tests and offline demonstration"),
        ("Postman/Swagger evidence", "Present in repository and report"),
        ("Frontend/backend/database screenshots", "Figures 1-24 embedded"),
        ("Atlas Free account and screenshots", "Complete; Figures 22-24 embedded"),
        ("Final merge and Assignment 03 tag", "Complete"),
        ("DOCX/PDF final export", "Complete"),
    ], widths=[3.8, 2.8], size=8.7)

    add_heading(doc, "Appendix A. Evidence File Index", 1)
    add_table(doc, ["Evidence group", "Files"], [
        ("Frontend and client persistence", "docs/screenshots/assignment-03/01-sign-in.png through 16-conflict-board.png"),
        ("Backend and database", "docs/screenshots/assignment-03/17-backend-task-stats.png through 21-backend-health-response.png"),
        ("Atlas Free deployment", "docs/screenshots/assignment-03/22-atlas-free-cluster.png through 24-atlas-connected-health.png"),
        ("Postman API collection", "docs/api-evidence/Huddle Assignment 02 API Evidence.postman_collection.json"),
        ("OpenAPI", "docs/openapi.yaml"),
        ("Verification scripts", "scripts/verify-assignment-02.mjs and scripts/verify-assignment-03.mjs"),
    ], widths=[2.3, 4.3], size=8.5)

    add_heading(doc, "Appendix B. Important Local Verification Commands", 1)
    add_code(doc, "npm ci\nnpm run seed\nnpm run dev:server\nnpm run dev\nnpm test\nnpm run build\nnpm run verify:assignment-02\nnpm run verify:assignment-03\nnpm audit --omit=optional")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
